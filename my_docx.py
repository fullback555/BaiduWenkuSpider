#!/usr/bin/env python
# -*- coding: utf-8 -*-
'''
@File       :  my_docx.py
@Contact    :  fullback555@163.com
@Modify Time:  2020/2/29 上午8:50 
@Author     :  pymark0202
@Desc       :  docx文件操作
'''
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import RGBColor, Pt, Inches, Cm  # inches:英寸，Cm:厘米
from docx.oxml.ns import qn
import os, sys

# document.styles['Normal'].font.name = u'宋体'
# document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
'''
1.最顶层是Document对象，其代表整个文档
2.block-level（块等级），段落是常见的块等级，换行符结尾算一个段落，表格、图片和标题均属于块对象；对于块对象属性，
常见有对齐（alignment）、缩进（indent）以及行间距（space）等等
3.inline-level（内联等级），其属于块等级中，run是常见的内联等级，一个块对象可由多个run组成，特别是通过run可由设置
不同属性样式；文字、句子、段落均可作为内联对象；对于内联对象属性，常见有字体、大小、对齐以及颜色等等
'''

class Mydocx():
	'''对docx进行简单的二次封装，方便操作'''

	def __init__(self, filename=None):
		self.doc = Document()  # 空文档

		if filename is not None:
			try:
				self.doc = Document(filename)
				print('加载文档成功')
			except Exception as e:
				self.doc = Document()  # 空文档
				print('加载文档失败，原因是：', e)

		self.para = self.doc.add_paragraph()  # 添加空段落
		self.run = self.para.add_run()  # run对象

	# 打开docx文档
	def open(self, filename):
		try:
			self.doc = Document(filename)
			return self.doc
		except Exception as e:
			print('打开文档失败，原因是：', e)

	# 保存docx文档
	def save(self, filename):
		self.doc.save(filename)
		return True

	# 读取docx
	def read(self, filename):
		doc = Document(filename)
		fullText = []
		for para in doc.paragraphs:
			fullText.append(para.text)
		return '\n'.join(fullText)

	# 添加标题
	def add_head(self, level=1):
		self.para = self.doc.add_heading(level=level)
		return self.para

	# 添加段落
	def add_para(self):
		self.para = self.doc.add_paragraph()
		return self.para

	# 添加run
	def add_run(self, text):
		self.run = self.para.add_run(text)
		return self.run

	# 添加标题，para为paragraph对象
	def add_text(self, text, family=u'宋体', size=-1, color=(125, 125, 125), bold=False, alignment='left'):
		# 添加文本
		self.run = self.para.add_run(text)
		# 设置字体  对于一些中文字体上述用font.name方法是无效的，需要使用_element.rPr.rFonts的set()方法
		self.run.font.name = family
		self.run._element.rPr.rFonts.set(qn('w:eastAsia'), family)
		# 设置字体大小
		if size != -1:
			self.run.font.size = Pt(size)
		# 设置字体颜色
		if color != (125, 125, 125):
			r, g, b = color
			self.run.font.color.rgb = RGBColor(r, g, b)
		# 设置文本加粗
		if bold:
			self.run.font.bold = True
		# 设置对齐方式
		if alignment in ('left', 'LEFT', 0):
			self.para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐
		elif alignment in ('center', 'CENTER', 1):
			self.para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中
		elif alignment in ('right', 'RIGHT', 2):
			self.para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 右对齐
		elif alignment in ('justify', 'JUSTIFY', 3):
			self.para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 完全两端对齐
		elif alignment in ('distribute', 'DISTRIBUTE', 4):
			self.para.alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE  # 段落字符被分布排列

		return self.run

	# 在一个段落前插入一新的段落
	def insert_before(self, para):
		new_para = para.insert_paragraph_before()
		self.para = new_para
		return self.para

	# 添加换行
	def add_enter(self):
		run = self.para.add_run()
		run.add_break()

	# 添加分页（换页）
	def add_page(self):
		self.doc.add_page_break()

	# 添加图片，一英寸等于2.54厘米
	def add_pic(self, filename, width=None):
		# 只指定宽度，会按图片比例自动设置高度。# inches:英寸，Cm:厘米
		if width is None:  # 不指定图片宽度
			self.doc.add_picture(filename)
		else:
			self.doc.add_picture(filename, width=Inches(width))
		# 或者在run上面添加图片（类似添加文本一样）
		# self.run.add_picture(filename, width=Inches(width))
		# self.doc.add_picture(filename, width=Inches(1), height=Cm(4))

	# 在run上添加图片，72dpi 1cm = 28.346像素
	def add_pic_in_run(self, filename, width=None):
		# 或者在run上面添加图片（类似添加文本一样）
		self.run = self.para.add_run()
		if width is None:  # 不指定图片宽度
			self.run.add_picture(filename)
		else:
			self.run.add_picture(filename, width=Cm(width))
		# self.doc.add_picture(filename, width=Inches(1), height=Cm(4))

	# 添加表格
	# def add_table(self, data, rows=2, cols=2, style="Normal Table"):
	def add_table(self, data, row_height=0.8, col_width=3, style="Table Grid"):
		if len(data) == 0:
			return False
		rows = len(data)
		cols = len(data[0])

		table = self.doc.add_table(rows=rows, cols=cols, style=style)
		for r in range(len(data)):
			for c in range(len(data[r])):
				table.cell(r, c).text = data[r][c]
				table.cell(r, c).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  #水平居中对齐
				table.cell(r, c).vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 垂直居中对齐
				# table.cell(r, c).vertical_alignment = WD_TABLE_ALIGNMENT.CENTER     # 垂直居中对齐
				# table.cell(r, c).horizontal_alignment = WD_TABLE_ALIGNMENT.CENTER   # 水平居中对齐
				table.rows[r].height = Cm(row_height)  # 设置行高
				table.cell(r, c).width = Cm(col_width)  # 设置列宽
			
		table.alignment = WD_TABLE_ALIGNMENT.CENTER     # 表格居中对齐
		return True


if __name__ == '__main__':

	# path = 'paper/A4.docx'
	mydocx = Mydocx()

	# 读取文件docx
	# filename = os.path.join(os.getcwd(), '清华.docx')
	# txt = mydocx.read(filename)
	# print(txt)

	print(mydocx.para, mydocx.run)
	mydocx.add_para()
	mydocx.add_text(text='我是段落第一行', alignment=1)
	print(mydocx.para.text, mydocx.run.text)

	# 添加标题
	p = mydocx.add_head(level=1)
	mydocx.add_text(text='我是一级标题', family=u'方正小标宋', size=20, color=(255, 0, 0), alignment='center')
	mydocx.add_head(level=3)
	mydocx.add_text(text='我是三级标题')

	# 添加段落
	p = mydocx.add_para()
	mydocx.add_text(text='我是段落一，你还好吗？', family='黑体', size=16, color=(0, 0, 255), alignment='left')
	mydocx.add_para()
	mydocx.add_text(text='\n我是段落二,我爱你。\n我有二行哦')

	# 添加段落
	p = mydocx.insert_before(mydocx.para)
	mydocx.add_text(text='我是段落二的上一个新的段落，插入进来的')

	mydocx.add_para()
	mydocx.add_text(text='我是段落三，右对齐哦。', alignment=2)

	# 添加换行
	mydocx.add_enter()
	mydocx.add_para()
	mydocx.add_text(text='我是段落四，我前面添加了换行', alignment=0)

	# 添加分页
	mydocx.add_page()
	mydocx.add_para()
	mydocx.add_text(text='我是段落五，我前面添加了分页')

	# 添加图片
	mydocx.add_pic_in_run(filename='girl.jpeg', width=10)
	mydocx.add_pic(filename='girl.jpeg', width=5)

	# 添加表格
	data = [['列1', '列1', '列1', '列1'],
			['11', '12', '13', '14'],
			['21', '22', '23', '24'],
			['31', '32', '33', '34']]
	mydocx.add_table(data)

	# 保存文档
	filename = os.path.join(os.path.dirname(sys.argv[0]), 'docxtest.docx')
	# filename = os.path.join(os.getcwd(), 'docxtest.docx')
	mydocx.save(filename)

